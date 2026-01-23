from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io


def generate_docx_file(data):
    """
    Generate a DOCX file from structured resume data.
    """
    doc = Document()

    # helper for styling
    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        return p

    # --- Header ---
    name = data.get("name", "Name Not Found")
    contact = data.get("contact_info", "")

    h1 = doc.add_heading(name, 0)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p_contact = doc.add_paragraph(contact)
    p_contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Summary ---
    if "summary" in data and data["summary"]:
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(data["summary"])

    # --- Experience ---
    if "experience" in data and data["experience"]:
        doc.add_heading("Experience", level=1)
        for job in data["experience"]:
            company = job.get("company", "")
            role = job.get("role", "")
            duration = job.get("duration", "")

            p = doc.add_paragraph()
            runner = p.add_run(f"{company} - {role}")
            runner.bold = True

            p.add_run(f"\n{duration} | {job.get('location', '')}")

            for point in job.get("points", []):
                doc.add_paragraph(point, style="List Bullet")

    # --- Education ---
    if "education" in data and data["education"]:
        doc.add_heading("Education", level=1)
        for edu in data["education"]:
            school = edu.get("institution", "")
            degree = edu.get("degree", "")
            year = edu.get("year", "")

            p = doc.add_paragraph()
            r = p.add_run(f"{school}")
            r.bold = True
            p.add_run(f"\n{degree} - {year}")

    # --- Skills ---
    if "skills" in data and data["skills"]:
        doc.add_heading("Skills", level=1)
        skills_text = ", ".join(data["skills"])
        doc.add_paragraph(skills_text)

    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream
